from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from app.schemas import DocumentFragment, ParsedDocument


def parse_docx(file_path: str) -> ParsedDocument:
    document = DocxDocument(file_path)
    fragments: list[DocumentFragment] = []
    current_section: str | None = None

    for order, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if "heading" in style_name or "заголов" in style_name:
            current_section = text
        fragments.append(
            DocumentFragment(
                text=text,
                page_from=None,
                page_to=None,
                section_title=current_section,
                order=order,
            )
        )

    return ParsedDocument(
        file_path=str(Path(file_path).resolve()),
        file_name=Path(file_path).name,
        file_type="docx",
        fragments=fragments,
    )
